import requests
from urllib.parse import quote # for double encoding IRIs for the API calls
import time
import json
from docx import Document

def retrieve_ontology_matches(term, numresults=50, ontology=None):
    """
    Retrieve ontology matches for a given term using OLS API

    params:
    - term (str) - a search term
    - ontologies (list) - a list of ontology names to search.
                            If not provided, all ontologies will be searched.   

    # TODO: smarter search (e.g., prioritize terms, use boolean logic)
    """
    base_url = 'https://www.ebi.ac.uk/ols4/api/search'
    # term = "alzheimer's"

    params = {
        'q': term,
        'rows': numresults,  
    }
    if ontology is not None:
        params['ontology'] = ontology

    try:
        # Make the API request
        response = requests.get(base_url, params=params)
        
        # Raise an exception for bad responses
        response.raise_for_status()
        
        # Parse the JSON results with full details
        results = response.json()

        jsonresults = json.dumps(results, indent=2)
        
        # Extract detailed concept information
        ontology_matches = []
        for term_info in results.get('response', {}).get('docs', []):
            # print(term_info)
            concept = {
                'label': term_info.get('label', 'No Label'),
                'iri': term_info.get('iri', 'No IRI'),
                'description': term_info.get('description', 'No Description'),
                'ontology_name': term_info.get('ontology_name', 'No Ontology Name'),
                'short_form': term_info.get('short_form', 'No Short Form')
            }
            ontology_matches.append(concept)

        return jsonresults, ontology_matches

    except requests.RequestException as e:
        print(f"Error connecting to OLS API: {e}")

def get_term_ancestors(ontology, iri):
    """
    Get ontology term properties using OLS API

    params:
    - ontology (str) - the ontology name
    - iri (str) - the IRI of the term

    returns:
    - ancestors (list) - list of ancestors, each with a dictionary of term properties

    Here is an example of a valid request:
    https://www.ebi.ac.uk/ols4/api/ontologies/duo/terms/http%253A%252F%252Fpurl.obolibrary.org%252Fobo%252FDUO_0000017/ancestors?lang=en
    """
    
    
    # double quote here if we're passing it directly into the URL
    # single quote if we're passing as param
    iri_encoded = quote(quote(iri, safe=''))
    
    # print(f"Attempting to get ancestors for {iri_encoded}")

    url = f"https://www.ebi.ac.uk/ols4/api/ontologies/{ontology}/terms/{iri_encoded}/ancestors"
    params = {
        'lang': 'en'
    }

    try:
        # Make the API request
        response = requests.get(url, params=params)
        
        # Raise an exception for bad responses
        response.raise_for_status()
        
        # Parse the JSON results with full details
        result = response.json()

        # get the ancestors
        ancestors = [i for i in result.get('_embedded', {}).get('terms', [])]
        
        return ancestors

    except requests.RequestException as e:
        
        print(f"Error connecting to OLS API: {e}")

        return e
    
def rank_ontology_matches_by_distance(ontology_matches):
    """
    Rank ontology matches by distance from root
    """
    # get the distance from root for each one
    # which is just the length of the ancestors list
    matches_with_metadata = []
    for idx, match in enumerate(ontology_matches):
        
        # get ancestors
        iri = match['iri']
        ontology = match['ontology_name']
        ancestors = get_term_ancestors(ontology, iri)
        print(f"Retrieved {len(ancestors)} ancestors for {iri}")

        # update match with ancestor info
        match['ancestors'] = ancestors
        match['distance_from_root'] = len(ancestors)

        matches_with_metadata.append(match)

        # pause for a bit
        # so we don't get rate limited
        time.sleep(DELAY)

    # sort by distance from root
    matches_with_metadata.sort(key=lambda x: x['distance_from_root'])
    
    return matches_with_metadata

def get_highlight_color(run):
    """Get the highlight color of a run."""
    if run.font.highlight_color:
        return run.font.highlight_color  # Word highlight color index
    elif run.font.color and run.font.color.rgb:
        return run.font.color.rgb  # RGB color
    return None

# TODO: grab the metadata for the document
def extract_peco_highlights_from_tables(doc_path):
    """Extract highlighted text from 'PECO statement' rows in tables."""
    doc = Document(doc_path)
    tables_data = []

    annotation_map = {
        'PINK (5)': 'Population',
        'BRIGHT_GREEN (4)': 'Exposure',
        'TURQUOISE (3)': 'Comparator',
        'YELLOW (7)': 'Outcome'
    }
    print(f"Found {len(doc.tables)} tables in the document")
    for table in doc.tables:
        table_dict = {}
        for row in table.rows:
            # doc title
            if "Title of manuscript" in row.cells[0].text.strip():
                table_dict["doc_title"] = row.cells[-1].text.strip()
            
            # doc name
            if "Last name of first author" in row.cells[0].text.strip():
                table_dict["doc_lastauthorname"] = row.cells[-1].text.strip()
            
            # doc year
            if "Year of publication" in row.cells[0].text.strip():
                table_dict["doc_year"] = row.cells[-1].text.strip()
            
            # url
            if "URL of HTML manuscript" in row.cells[0].text.strip():    
                table_dict["doc_url"] = row.cells[-1].text.strip()

            # section
            if "Section PECO statement is in" in row.cells[0].text.strip():
                table_dict["doc_peco_section"] = row.cells[-1].text.strip()

            # annotation note
            if "Annotator comments" in row.cells[0].text.strip():
                table_dict["doc_annotator_comments"] = row.cells[-1].text.strip()

            # PECO statement
            if "PECO statement" in row.cells[0].text.strip():  # Check for PECO row
                peco_cell = row.cells[-1]  # Assume PECO text is in the 2nd column
                color_text_map = {}

                statement = ""
                for para in peco_cell.paragraphs:
                    for run in para.runs:
                        color = get_highlight_color(run)
                        text = run.text.strip()
                        if color and text:  # Only include highlighted text
                            color_key = str(color)
                            # print(f"Found highlighted text: {text} with color {color_key}")
                            if color_key in annotation_map:
                                color_key = annotation_map[color_key]
                                if color_key not in color_text_map:
                                    color_text_map[color_key] = []
                                color_text_map[color_key].append(text)
                        statement = statement + "" + text + " "

                table_dict["peco_elements"] = color_text_map
                table_dict['peco_statement'] = statement
        
        if table_dict:
            tables_data.append(table_dict)

    return tables_data

NUMRESULTS = 5
DELAY = 0.05

doc_path = "data/PECO-examples_merged-26.docx"  

# Extract PECO highlights from tables in the document
annotations = extract_peco_highlights_from_tables(doc_path)
print(f"Extracted {len(annotations)} PECO statement highlights in the document")

parsed_output = []

# for each highlight, retrieve ontology matches
# TODO: handle keyphrases
for annotation in annotations:
    print(f"Retrieving ontology matches for {annotation['doc_title']}")
    print(f"PECO elements: {annotation['peco_elements']}")
    annotation_d = {
        "doc_title": annotation["doc_title"],
        "doc_lastauthorname": annotation["doc_lastauthorname"],
        "doc_year": annotation["doc_year"],
        "url": annotation["doc_url"],
        "doc_peco_section": annotation["doc_peco_section"],
        'peco_statement': annotation['peco_statement'],
        "peco_elements": {
            'Population': [],
            'Exposure': [],
            'Comparator': [],
            'Outcome': []
        }
    }
    for peco_element, el_texts in annotation["peco_elements"].items():
        for text in el_texts:
            # get ontology matches
            jsonresults, ontology_matches = retrieve_ontology_matches(text, numresults=NUMRESULTS)
            # get ancestors and rank by distance from root
            # removing this for now just to speed up the process
            # ranked_matches = rank_ontology_matches_by_distance(ontology_matches)

            # get current list of elements for peco element
            current_elements = annotation_d["peco_elements"][peco_element]
            current_elements.append({
                "text": text,
                # "matches": ranked_matches
                "matches": ontology_matches
            })
    parsed_output.append(annotation_d)

with open("data/parsed_output.json", "w") as f:
    json.dump(parsed_output, f, indent=2)

 
# # Print the structured output
# for i, table in enumerate(result, 1):
#     print(f"Table {i}:")
#     for color, texts in table["PECO statement highlights"].items():
#         print(f"  Color: {color}, Text: {', '.join(texts)}")
#     print("-" * 50)

# html = """

# """
# # Define the HTML table structure
# html_table = """
# <table border="1">
#   <tr>
#     <th>Preferred Term</th>
#     <th>Ontology</th>
#   </tr>
# """

# # Loop through the data and add each row to the HTML table
# for row in parsed_output:
#     html_table += """
#   <tr>
#     <td>{Name}</td>
#     <td>{Age}</td>
#     <td>{City}</td>
#   </tr>
# """.format(**row)

# # Close the table tag
# html_table += "</table>"

# # Print the HTML table to the console
# print(html_table)